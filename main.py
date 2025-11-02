from fastmcp import FastMCP
import re
from pathlib import Path
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from pathlib import Path
import imaplib

# Initialize FastMCP server
mcp = FastMCP("Cold Email Assistant")


@mcp.tool()
def send_email_tool(
    receiver_email: str,
    subject: str,
    body: str,
    sender_password: str = 'klsk wrvf bwzq azez',
    sender_email: str = 'ameymedewar001@gmail.com',
    resume_path: str = "resume.pdf"
) -> str:
    """
    MCP tool to send an email with subject, body, and attach a resume (.docx).
    """
    try:
        msg = MIMEMultipart()
        msg["From"] = sender_email
        msg["To"] = receiver_email
        msg["Subject"] = subject
        msg.attach(MIMEText(body, "plain"))

        resume_file = Path(resume_path)
        if not resume_file.exists():
            return f"Error: Resume file '{resume_file}' not found."

        with open(resume_file, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())

        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f"attachment; filename={resume_file.name}")
        msg.attach(part)

        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)

        return f"Email sent successfully to {receiver_email}"

    except Exception as e:
        return f"Error sending email: {str(e)}"


@mcp.tool()
def save_to_draft_tool(
    receiver_email: str,
    subject: str,
    body: str,
    sender_password: str = 'klsk wrvf bwzq azez',
    sender_email: str = 'ameymedewar001@gmail.com',
    resume_path: str = "resume.pdf"
) -> str:
    """
    MCP tool to save an email as a draft in Gmail using IMAP.
    Uses the same App Password as send_email_tool.
    """
    try:
        import imaplib
        from time import time
        from email.utils import formatdate
        
        # Create message
        msg = MIMEMultipart()
        msg["From"] = sender_email
        msg["To"] = receiver_email
        msg["Subject"] = subject
        msg["Date"] = formatdate(localtime=True)
        msg.attach(MIMEText(body, "plain"))

        # Attach resume
        resume_file = Path(resume_path)
        if not resume_file.exists():
            return f"Error: Resume file '{resume_file}' not found."

        with open(resume_file, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())

        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f"attachment; filename={resume_file.name}")
        msg.attach(part)

        # Connect to Gmail IMAP
        imap = imaplib.IMAP4_SSL("imap.gmail.com")
        imap.login(sender_email, sender_password)
        
        # Append to Drafts folder with proper flags
        imap.append('[Gmail]/Drafts', '\\Draft', imaplib.Time2Internaldate(time()), msg.as_bytes())
        imap.logout()

        return f"Draft saved successfully!\nRecipient: {receiver_email}\nSubject: {subject}\n\nCheck your Gmail Drafts folder."

    except Exception as e:
        return f"Error saving draft: {str(e)}"


@mcp.tool()
def parse_job_message(message: str) -> dict:
    """
    Parse WhatsApp job message to extract company name, role, email, location, 
    stipend, and other job details.
    
    Args:
        message: The WhatsApp message text containing job posting details
        
    Returns:
        Dictionary with extracted job details: company_name, role, email, 
        location, stipend, batch, requirements, raw_message
    """
    details = {
        "company_name": None,
        "role": None,
        "email": None,
        "location": None,
        "stipend": None,
        "batch": None,
        "requirements": None,
        "raw_message": message
    }
    
    # Extract email
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, message)
    if emails:
        details["email"] = emails[0]
    
    # Extract company name
    company_match = re.search(r'Company\s*[-:]\s*(.+?)(?:\n|$)', message, re.IGNORECASE)
    if company_match:
        details["company_name"] = company_match.group(1).strip()
    
    # Extract role
    role_match = re.search(r'Role\s*[-:]\s*(.+?)(?:\n|$)', message, re.IGNORECASE)
    if role_match:
        details["role"] = role_match.group(1).strip()
    
    # Extract location
    location_match = re.search(r'Location\s*[-:]\s*(.+?)(?:\n|$)', message, re.IGNORECASE)
    if location_match:
        details["location"] = location_match.group(1).strip()
    
    # Extract stipend/salary
    stipend_match = re.search(r'Stipend\s*[-:]\s*(.+?)(?:\n|$)', message, re.IGNORECASE)
    if stipend_match:
        details["stipend"] = stipend_match.group(1).strip()
    
    # Extract batch/graduation year
    batch_match = re.search(r'Batch\s*[-:]\s*(.+?)(?:\n|$)', message, re.IGNORECASE)
    if batch_match:
        details["batch"] = batch_match.group(1).strip()
    
    # Extract requirements section
    req_match = re.search(r'Requirements?:(.+?)(?=How to Apply:|$)', message, re.IGNORECASE | re.DOTALL)
    if req_match:
        requirements = req_match.group(1).strip()
        # Clean up bullet points
        requirements = re.sub(r'[•\-]\s*', '- ', requirements)
        details["requirements"] = requirements
    
    return details

@mcp.tool()
def load_resume(filename: str = "resume.docx") -> str:
    """
    Load resume content from a Word document (.docx file) in the project directory.
    
    Args:
        filename: Name of the resume file (default: resume.docx)
        
    Returns:
        String containing the full text content of the resume
    """
    try:
        # Get the path to the resume file
        resume_path = Path(__file__).parent / filename
        
        if not resume_path.exists():
            return f"Error: Resume file '{filename}' not found in {resume_path.parent}"
        
        # Read the .docx file
        doc = Document(resume_path)
        
        # Extract all text from paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text)
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        resume_content = '\n'.join(full_text)
        
        if not resume_content.strip():
            return "Error: Resume file is empty"
        
        return resume_content
    
    except Exception as e:
        return f"Error loading resume: {str(e)}"

@mcp.tool()
def load_prompt(filename: str = "prompt.docx") -> str:
    """
    Load prompt content from a Word document (.docx file) in the project directory.
    
    Args:
        filename: Name of the prompt file (default: prompt.docx)
        
    Returns:
        String containing the full text content of the prompt
    """
    try:
        # Get the path to the prompt file
        prompt_path = Path(__file__).parent / filename
        
        if not prompt_path.exists():
            return f"Error: prompt file '{filename}' not found in {prompt_path.parent}"
        
        # Read the .docx file
        doc = Document(prompt_path)
        
        # Extract all text from paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text)
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        prompt_content = '\n'.join(full_text)
        
        if not prompt_content.strip():
            return "Error: prompt file is empty"
        
        return prompt_content
    
    except Exception as e:
        return f"Error loading prompt: {str(e)}"

    
if __name__ == "__main__":
    mcp.run()